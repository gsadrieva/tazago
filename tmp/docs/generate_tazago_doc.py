from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Document\Build8")
OUTPUT = ROOT / "output" / "doc" / "TazaGo_Tehnikalyk_Quzhat_v2_KK.docx"


ACCENT = RGBColor(24, 76, 146)
TEXT = RGBColor(34, 34, 34)
MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_page_margins(section) -> None:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT

    for style_name, size in [("Title", 21), ("Heading 1", 15), ("Heading 2", 12)]:
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.color.rgb = ACCENT
        style.font.size = Pt(size)


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TazaGo жобасының техникалық құжаты")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = TEXT

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Жоба құрылымы, негізгі файлдар мен техникалық база туралы қысқаша анықтама")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED

    document.add_paragraph("")
    info = [
        "Жоба: TazaGo",
        "Құжат тілі: қазақ тілі",
        "Құжат түрі: ішкі техникалық құжаттама",
        "Мақсаты: код базасының құрылымын тез түсіндіру",
    ]
    for line in info:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    document.add_page_break()


def add_summary(document: Document) -> None:
    document.add_heading("1. Жоба туралы қысқаша мәлімет", level=1)
    intro = (
        "TazaGo — cleaning-сервиске арналған веб-платформа. Жоба клиентке қызметтерді көруге, "
        "тапсырыс беруге, WhatsApp арқылы нақтылауға, ал әкімшіге тапсырыстарды басқаруға мүмкіндік береді."
    )
    document.add_paragraph(intro)

    bullets = [
        "Фронтенд: Next.js 16, React 19, TypeScript.",
        "Стильдеу: Tailwind CSS 4 және globals.css ішіндегі жоба стилдері.",
        "Бэкенд: Supabase Auth, Postgres, RPC функциялары.",
        "Тілдік қолдау: қазақша және орысша.",
        "Негізгі модульдер: каталог, booking wizard, dashboard, admin-панель.",
    ]
    for item in bullets:
        document.add_paragraph(item, style="List Bullet")


def add_architecture(document: Document) -> None:
    document.add_heading("2. Архитектура логикасы", level=1)
    rows = [
        ("Маршрут қабаты", "app/ бумасында орналасқан және Next.js App Router арқылы жұмыс істейді."),
        ("UI компоненттері", "components/ ішінде сақталады, беттер соларды жинап пайдаланады."),
        ("Бизнес-логика", "lib/ ішінде бөлінген: auth, actions, queries, i18n, format."),
        ("Дерекқор", "Supabase миграциялары мен RPC функциялары арқылы басқарылады."),
        ("Тапсырыс ағыны", "Сайтта форма толтырылады, содан кейін WhatsApp redirect жасалады."),
        ("Қауіпсіздік", "Supabase RLS саясаттары және бөлек admin-session механизмі қолданылады."),
    ]

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Қабат"
    header[1].text = "Міндеті"
    for cell in header:
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_folder_table(document: Document, title: str, rows: list[tuple[str, str]]) -> None:
    document.add_heading(title, level=1)
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    header[0].text = "Файл немесе бума"
    header[1].text = "Не үшін жауап береді"
    for cell in header:
        set_cell_shading(cell, "D9EAF7")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for name, description in rows:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = description
        row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_document() -> None:
    document = Document()
    set_page_margins(document.sections[0])
    configure_styles(document)
    add_title_page(document)
    add_summary(document)
    add_architecture(document)

    add_folder_table(
        document,
        "3. Жоғарғы деңгейдегі құрылым",
        [
            ("app/", "Негізгі беттер мен маршруттар орналасқан басты қолданба бумасы."),
            ("components/", "Қайта қолданылатын UI компоненттері."),
            ("lib/", "Логика, көмекші функциялар және Supabase интеграциясы."),
            ("public/", "Статикалық файлдар мен логотиптер."),
            ("supabase/", "Дерекқор миграциялары мен SQL логикасы."),
            ("types/", "TypeScript типтері мен интерфейстері."),
            (".env.example", "Қажетті environment айнымалыларының үлгісі."),
            ("package.json", "npm тәуелділіктері мен іске қосу скрипттері."),
            ("tsconfig.json", "TypeScript компиляциясының баптаулары."),
            ("proxy.ts", "Сұраныс кезінде Supabase сессиясын жаңартатын middleware."),
        ],
    )

    add_folder_table(
        document,
        "4. app/ бумасының негізгі файлдары",
        [
            ("app/layout.tsx", "Жобаның түпкі layout файлы, глобал қаріптер мен HTML қаңқасын береді."),
            ("app/page.tsx", "Түбір маршрутты әдепкі тілге бағыттайды."),
            ("app/globals.css", "Глобал стильдер, түстер, батырмалар және layout ережелері."),
            ("app/api/bookings/route.ts", "Booking API маршруты: сұранысты қабылдайды, тексереді, RPC шақырады."),
            ("app/auth/confirm/route.ts", "Auth растау сценарийіне арналған серверлік маршрут."),
            ("app/auth/signout/route.ts", "Пайдаланушыны жүйеден шығару маршруты."),
            ("app/[locale]/layout.tsx", "Белгілі бір тіл үшін metadata және route деңгейін басқарады."),
            ("app/[locale]/page.tsx", "Таңдалған тілдегі басты бет."),
            ("app/[locale]/services/page.tsx", "Қызметтер тізімі."),
            ("app/[locale]/services/[slug]/page.tsx", "Белгілі қызметтің толық беті."),
            ("app/[locale]/booking/page.tsx", "Тапсырыс беру беті."),
            ("app/[locale]/dashboard/page.tsx", "Клиенттің жеке кабинеті."),
            ("app/[locale]/admin/page.tsx", "Әкімшілік панель."),
            ("app/[locale]/admin/login/page.tsx", "Әкімшілік кіру беті."),
            ("app/[locale]/auth/sign-in/page.tsx", "Пайдаланушының кіру беті."),
            ("app/[locale]/auth/sign-up/page.tsx", "Тіркелу беті."),
            ("app/[locale]/contact/page.tsx", "Байланыс және FAQ беті."),
        ],
    )

    add_folder_table(
        document,
        "5. components/ бумасының негізгі файлдары",
        [
            ("components/site-shell.tsx", "Header, navigation, footer және жалпы сыртқы қаңқаны береді."),
            ("components/locale-switcher.tsx", "Қазақша/орысша тіл ауыстырғыш."),
            ("components/booking/booking-wizard.tsx", "Көпқадамды тапсырыс формасы және есептеу логикасы."),
            ("components/home/chat-widget.tsx", "Жиі сұрақтарға жауап беретін чат-виджет."),
            ("components/home/testimonials-carousel.tsx", "Клиент пікірлерінің каруселі."),
            ("components/ui/status-badge.tsx", "Статустарды визуалды badge түрінде көрсетеді."),
        ],
    )

    add_folder_table(
        document,
        "6. lib/ бумасының негізгі файлдары",
        [
            ("lib/auth.ts", "Сессияны алу, міндетті кіру және admin тексеру функциялары."),
            ("lib/admin-session.ts", "Cookie-негізді әкімшілік сессия логикасы."),
            ("lib/booking.ts", "Booking payload үшін Zod валидация схемасы."),
            ("lib/format.ts", "Баға мен күнді форматтау функциялары."),
            ("lib/i18n.ts", "Барлық аудармалар мен мәтіндік сөздіктер."),
            ("lib/locale.ts", "Тілді анықтау және locale helper-функциялары."),
            ("lib/actions/auth.ts", "Кіру мен тіркелу server action-дары."),
            ("lib/actions/admin-auth.ts", "Әкімшілік кіру/шығу action-дары."),
            ("lib/actions/admin.ts", "Admin тарапынан тапсырыстар мен қызметтерді басқару action-дары."),
            ("lib/queries/catalog.ts", "Қызметтер каталогын Supabase-тен жинайды."),
            ("lib/queries/bookings.ts", "Пайдаланушы және admin тапсырыстарын оқиды."),
            ("lib/supabase/server.ts", "Сервер жағындағы Supabase клиенті."),
            ("lib/supabase/client.ts", "Клиент жағындағы Supabase клиенті."),
            ("lib/supabase/proxy.ts", "Middleware ішінде сессия cookie-лерін синхрондайды."),
            ("lib/supabase/admin.ts", "Service role кілтімен әкімшілік Supabase клиенті."),
        ],
    )

    add_folder_table(
        document,
        "7. Дерекқор миграциялары",
        [
            ("202604080001_tazaqala_mvp.sql", "Негізгі кестелерді, trigger-лерді және RLS саясаттарын жасайды."),
            ("202604080002_tazaqala_booking_rpc.sql", "Тапсырыс құру және WhatsApp статусы үшін RPC функциялары."),
            ("202604080003_expand_addons.sql", "Қосымша қызметтерді кеңейтетін деректер миграциясы."),
            ("202604110001_admin_create_service.sql", "Admin-панель арқылы жаңа қызмет енгізуге арналған RPC."),
        ],
    )

    document.add_heading("8. Environment айнымалылары", level=1)
    for item in [
        "NEXT_PUBLIC_SUPABASE_URL - Supabase жобасының URL мекенжайы.",
        "NEXT_PUBLIC_SUPABASE_PUBLISHABLE_KEY - браузер жағында қолданылатын кілт.",
        "SUPABASE_SERVICE_ROLE_KEY - қорғалған серверлік операцияларға арналған кілт.",
        "NEXT_PUBLIC_SITE_URL - сайттың базалық URL мекенжайы.",
        "NEXT_PUBLIC_WHATSAPP_PHONE - booking redirect жасалатын WhatsApp нөмірі.",
        "ADMIN_LOGIN, ADMIN_PASSWORD, ADMIN_SESSION_SECRET - әкімшілік cookie-сессия параметрлері.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("9. Қорытынды", level=1)
    document.add_paragraph(
        "Код базасы логикалық түрде жақсы бөлінген: беттер мен маршруттар app/ ішінде, "
        "компоненттер components/ ішінде, ал бизнес-логика мен интеграциялар lib/ ішінде орналасқан. "
        "Supabase миграциялары дерекқор құрылымын бөлек, түсінікті қабат ретінде ұстап тұр. "
        "Бұл жобаға жаңа функция қосу, админ-панельді кеңейту немесе booking ағынын дамыту салыстырмалы түрде ыңғайлы."
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)
